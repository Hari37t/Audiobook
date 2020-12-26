import pyttsx3
import PyPDF2
import docx


def pdfread(filename):

    book = open(filename, 'rb')
    pdfReader = PyPDF2.PdfFileReader(book)
    pages = pdfReader.numPages
    print(pages)
    pagenumber=int(input("Which page do you want to read : "))
    print((type(pagenumber)))
    page = pdfReader.getPage(pagenumber)
    text = page.extractText()
    return (text)



def wordread(filename):

    wordfile = docx.Document(filename)
    print(len(wordfile.paragraphs))
    paranumber=int(input("Which paragraph do you want to read : "))
    text = wordfile.paragraphs[paranumber].text
    return text



def speak(text):
    speaker = pyttsx3.init()
    print(text)
    speaker.say(text)
    speaker.runAndWait()


#text =wordread("sainapse.docx")
#print(text)
#speak(text)

text =pdfread("C:\\Users\\Harihara.Manigandan\\Downloads\\progit.pdf")
print(text)
speak(text)




